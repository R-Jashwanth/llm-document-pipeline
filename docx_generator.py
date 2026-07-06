"""
Microsoft Word (.docx) Document Generator Module.

This module converts Markdown text (produced by the generation and reflection modules)
into a structured, professional Microsoft Word document using the python-docx SDK.

Design Decisions:
1. AST-style Parsing State Machine: Scans the markdown text line-by-line, parsing
   hierarchical headers, lists, and formatting tables dynamically.
2. Inline Formatting: Parses basic markdown bold (`**text**`) and italic (`*text*`) tags
   and applies them as styled Runs in python-docx paragraphs.
3. Tabular Layouts: Recognizes markdown pipe tables (`| Col1 | Col2 |`), strips divider
   rows, normalizes column counts, and maps them to a formatted Word table using 'Table Grid'.
"""

import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config import settings
from utils import DocxException, logger


class DocxGenerator:
    """
    DocxGenerator handles parsing markdown structures and compiling them into
    well-formatted MS Word documents saved in the output directory.
    """

    def create_docx(self, markdown_content: str, filename: str) -> Path:
        """
        Parses a markdown string and writes it to a styled Word Document.

        Args:
            markdown_content: The full markdown text to convert.
            filename: The target filename (e.g. 'project_proposal.docx').

        Returns:
            The Path where the generated document was saved.

        Raises:
            DocxException: If writing or saving the file fails.
        """
        logger.info(f"Starting DOCX compiler for filename: '{filename}'")
        doc = Document()

        # Define default typography style (11pt Calibri)
        normal_style = doc.styles["Normal"]
        font = normal_style.font
        font.name = "Calibri"
        font.size = Pt(11)

        lines = markdown_content.splitlines()
        num_lines = len(lines)
        i = 0

        try:
            while i < num_lines:
                line = lines[i].strip()

                # Skip empty lines
                if not line:
                    i += 1
                    continue

                # 1. Handle Headings (#, ##, ###)
                if line.startswith("#"):
                    level = 0
                    while level < len(line) and line[level] == "#":
                        level += 1

                    heading_text = line[level:].strip()

                    # Restrict level to a max of 3 for styling consistency
                    heading_level = min(level, 3)

                    # Add heading and configure alignment
                    h = doc.add_heading(heading_text, level=heading_level)
                    if heading_level == 1:
                        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    i += 1
                    continue

                # 2. Handle Tables (| Col1 | Col2 |)
                if line.startswith("|"):
                    table_lines = []
                    while i < num_lines and lines[i].strip().startswith("|"):
                        table_lines.append(lines[i].strip())
                        i += 1
                    self._add_table_to_doc(doc, table_lines)
                    continue

                # 3. Handle Bullet Lists (- item or * item)
                if line.startswith("- ") or line.startswith("* "):
                    list_text = line[2:].strip()
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_formatted_text(p, list_text)
                    i += 1
                    continue

                # 4. Handle Numbered Lists (1. item)
                num_list_match = re.match(r"^\d+\.\s+(.*)", line)
                if num_list_match:
                    list_text = num_list_match.group(1).strip()
                    p = doc.add_paragraph(style="List Number")
                    self._add_formatted_text(p, list_text)
                    i += 1
                    continue

                # 5. Handle Regular Paragraphs
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
                i += 1

            # Resolve full target save path
            output_path = settings.generated_dir / filename
            doc.save(str(output_path))
            logger.info(f"DOCX document successfully written to: '{output_path}'")
            return output_path

        except Exception as e:
            logger.error(f"Error compiling Word document: {e}")
            raise DocxException(
                "Failed to convert Markdown content to DOCX format.", details=str(e)
            ) from e

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """
        Parses basic inline markdown formatting (**bold** and *italic*)
        and appends runs to the given paragraph.
        """
        # Regex splits on bold and italic tokens
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for token in tokens:
            if not token:
                continue

            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)

    def _add_table_to_doc(self, doc: Document, table_lines: list[str]) -> None:
        """
        Parses raw markdown table lines, cleans delimiters, and constructs
        a native docx Table styled with borders and bold headers.
        """
        if len(table_lines) < 2:
            return  # Needs at least a header row and content/divider row

        # Parse header row
        header_raw = table_lines[0]

        # Check if the second row is a Markdown divider row (e.g. |---|---|)
        data_start_idx = 1
        if len(table_lines) > 1 and re.match(r"^[\s|:-]+$", table_lines[1]):
            data_start_idx = 2

        # Parser helper for rows split by pipes
        def parse_row(row_str: str) -> list[str]:
            cells = [cell.strip() for cell in row_str.split("|")]
            # Strip empty elements created by outer pipes
            if row_str.startswith("|") and len(cells) > 0:
                cells = cells[1:]
            if row_str.endswith("|") and len(cells) > 0:
                cells = cells[:-1]
            return cells

        headers = parse_row(header_raw)
        num_cols = len(headers)
        if num_cols == 0:
            return

        # Parse subsequent rows
        data_rows = []
        for line in table_lines[data_start_idx:]:
            row_data = parse_row(line)
            # Normalize column count (pad or truncate)
            if len(row_data) < num_cols:
                row_data += [""] * (num_cols - len(row_data))
            elif len(row_data) > num_cols:
                row_data = row_data[:num_cols]
            data_rows.append(row_data)

        # Build table structure
        total_rows = 1 + len(data_rows)
        table = doc.add_table(rows=total_rows, cols=num_cols)
        table.style = "Table Grid"

        # Populate header cells and style as bold
        hdr_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(headers):
            hdr_cells[col_idx].text = col_name
            # Apply bold styling to all paragraphs/runs in header cells
            for paragraph in hdr_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Populate data rows
        for r_idx, r_data in enumerate(data_rows):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, cell_value in enumerate(r_data):
                # We can add text and check for inline formatting
                paragraph = row_cells[c_idx].paragraphs[0]
                self._add_formatted_text(paragraph, cell_value)
